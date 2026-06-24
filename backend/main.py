import io
import os
import re
import time
import uuid
import base64
import tempfile
from urllib.parse import quote
from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.responses import StreamingResponse
from fastapi.middleware.cors import CORSMiddleware
import pypandoc
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from base_model.markdown_request import MarkdownRequest

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
    expose_headers=["Content-Disposition"]
)

in_memory_store = {}


def add_field_run(paragraph, instruction: str, placeholder_text: str = "") -> None:
    run = paragraph.add_run()

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")

    instruction_text = OxmlElement("w:instrText")
    instruction_text.set(qn("xml:space"), "preserve")
    instruction_text.text = instruction

    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")

    text = OxmlElement("w:t")
    text.text = placeholder_text

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    run._r.append(begin)
    run._r.append(instruction_text)
    run._r.append(separate)
    run._r.append(text)
    run._r.append(end)

def enable_field_updates(document: Document) -> None:
    settings_element = document.settings.element
    update_fields = settings_element.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings_element.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def add_page_numbers(document: Document) -> None:
    for section in document.sections:
        footer = section.footer
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        if paragraph.text.strip():
            paragraph = footer.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_field_run(paragraph, "PAGE", "1")


def add_table_of_contents(document: Document) -> None:
    heading_paragraph = document.add_paragraph("Table of Contents")
    toc_paragraph = document.add_paragraph()
    add_field_run(
        toc_paragraph,
        'TOC \\o "1-3" \\h \\z \\u',
        "Right-click to update the table of contents.",
    )
    body = document._body._element
    body.insert(0, toc_paragraph._p)
    body.insert(0, heading_paragraph._p)


def customize_docx_output(docx_path: str, request: MarkdownRequest) -> None:
    if not request.settings.includeToc and not request.settings.includePageNumbers:
        return
    document = Document(docx_path)
    enable_field_updates(document)
    if request.settings.includeToc:
        add_table_of_contents(document)
    if request.settings.includePageNumbers:
        add_page_numbers(document)
    document.save(docx_path)


def process_mermaid_blocks(md_content: str, tmp_dir: str) -> str:
    pattern = r"!\[diagram\]\((data:image/png;base64,([^)]+))\)"

    def replacer(match):
        base64_data = match.group(2)
        img_filename = f"{uuid.uuid4()}.png"
        img_path = os.path.join(tmp_dir, img_filename)

        with open(img_path, "wb") as f:
            f.write(base64.b64decode(base64_data))

        return f"![diagram]({img_filename})"

    return re.sub(pattern, replacer, md_content)


def process_and_store_conversion(request: MarkdownRequest) -> dict:
    base_filename = request.settings.filename
    md_content = request.content

    if base_filename.endswith('.md'):
        base_filename = base_filename.rsplit(".", 1)[0]
    final_filename = base_filename if base_filename.endswith(".docx") else f"{base_filename}.docx"

    temp_output = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    temp_output_path = temp_output.name
    temp_output.close()

    try:
        with tempfile.TemporaryDirectory() as tmpdir:
            # decode base64 PNG diagrams to files
            md_content = process_mermaid_blocks(md_content, tmpdir)

            md_input_path = os.path.join(tmpdir, "input.md")
            with open(md_input_path, "w") as f:
                f.write(md_content)

            docx_template = request.settings.template

            pypandoc.convert_file(
                md_input_path,
                'docx',
                extra_args=[
                    '--standalone',
                    f'--reference-doc={docx_template}.docx',
                    f'--resource-path={tmpdir}'
                ],
                outputfile=temp_output_path
            )

        if not os.path.exists(temp_output_path) or os.path.getsize(temp_output_path) == 0:
            raise RuntimeError("Pandoc produced empty output")

        customize_docx_output(temp_output_path, request)

        with open(temp_output_path, "rb") as f:
            docx_bytes = f.read()

        buffer = io.BytesIO(docx_bytes)
        file_id = str(uuid.uuid4())
        size = buffer.getbuffer().nbytes

        in_memory_store[file_id] = {
            "filename": final_filename,
            "content": buffer,
            "size": size,
            "stored_at": time.time()
        }

        return {
            "file_id": file_id,
            "filename": final_filename,
            "size": size
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Conversion failed: {str(e)}")
    finally:
        if os.path.exists(temp_output_path):
            os.remove(temp_output_path)


@app.post("/convert-text/")
async def convert_text_to_docx(request: MarkdownRequest) -> dict:
    if not request.settings.filename:
        raise HTTPException(status_code=400, detail="Filename is required")
    return process_and_store_conversion(request)


@app.post("/convert/")
async def convert_md_to_docx(file: UploadFile = File(...)) -> dict:
    if not file.filename or not file.filename.endswith('.md'):
        raise HTTPException(status_code=400, detail="Only .md files are allowed")
    content = await file.read()
    try:
        decoded = content.decode("utf-8")
    except UnicodeDecodeError:
        raise HTTPException(status_code=400, detail="File is not valid UTF-8")
    request = MarkdownRequest(
        content=decoded,
        settings={"filename": file.filename}
    )
    return process_and_store_conversion(request)


@app.get("/download/{file_id}")
async def download_file(file_id: str) -> StreamingResponse:
    now = time.time()
    expired = [fid for fid, data in in_memory_store.items()
               if now - data.get("stored_at", 0) > 3600]
    for fid in expired:
        del in_memory_store[fid]

    file_data = in_memory_store.get(file_id)
    if not file_data:
        raise HTTPException(status_code=404, detail="File not found or expired")

    buffer = file_data["content"]
    filename = file_data["filename"]
    buffer.seek(0)

    headers = {
        "Content-Disposition": f"attachment; filename*=UTF-8''{quote(filename)}",
        "Access-Control-Expose-Headers": "Content-Disposition"
    }

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers
    )